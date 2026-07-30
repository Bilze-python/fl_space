from __future__ import annotations

import json
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch


ROOT = Path(r"D:\fl_space")
SOURCE_DOCX_PATH = ROOT / "SpaceFL_人机协作论文报告.docx"
OUTPUT_DOCX_PATH = ROOT / "SpaceFL_人机协作论文报告_完善版.docx"
BACKUP_PATH = ROOT / "SpaceFL_人机协作论文报告.backup_before_enhance.docx"
GRID_SUMMARY_PATH = ROOT / "experiment_output" / "grid_summary.json"
FEDPROX_SUMMARY_PATH = ROOT / "results" / "experiment_summary.json"
GENERATED_DIR = ROOT / "docs" / "generated_report_assets"
FRAMEWORK_PNG = GENERATED_DIR / "spacefl_framework_cn.png"
SUMMARY_PNG = GENERATED_DIR / "spacefl_experiment_summary_cn.png"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def pick_font() -> str:
    candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Arial Unicode MS"]
    return candidates[0]


def configure_matplotlib() -> None:
    plt.rcParams["font.sans-serif"] = [
        "Microsoft YaHei",
        "SimHei",
        "Noto Sans CJK SC",
        "Arial Unicode MS",
        "DejaVu Sans",
    ]
    plt.rcParams["axes.unicode_minus"] = False


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None):
    font_name = pick_font()
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_cn_paragraph(doc: Document, text: str, *, style: str = "Normal", bold_prefix: str | None = None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = para.add_run(bold_prefix)
        set_run_font(prefix_run, size=11, bold=True, color="1F1F1F")
        rest_run = para.add_run(text[len(bold_prefix):])
        set_run_font(rest_run, size=11, color="1F1F1F")
    else:
        run = para.add_run(text)
        set_run_font(run, size=11, color="1F1F1F")
    return para


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    set_run_font(run, size=10, color="4F4F4F")
    run.italic = True


def build_framework_figure(path: Path) -> None:
    configure_matplotlib()
    fig, ax = plt.subplots(figsize=(14, 8), dpi=220)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    title = "SpaceFL 技术框架与人机协同质控闭环"
    ax.text(0.5, 0.95, title, ha="center", va="center", fontsize=20, fontweight="bold", color="#103B2A")

    boxes = [
        (
            0.05,
            0.58,
            0.16,
            0.22,
            "#EAF4EF",
            "环境层\nEnvironment",
            "天体参数\n大气模型\n地面站网络",
        ),
        (
            0.24,
            0.58,
            0.16,
            0.22,
            "#E3F0FB",
            "轨道层\nOrbit",
            "Kepler 双体\nSkyfield/SGP4\n可见性计算",
        ),
        (
            0.43,
            0.58,
            0.16,
            0.22,
            "#F3F1FF",
            "模拟层\nSimulator",
            "接触矩阵\n时隙推进\nISL 可选",
        ),
        (
            0.62,
            0.58,
            0.16,
            0.22,
            "#EEF7E7",
            "联邦学习层\nFL",
            "FedAvg / FedProx / FedBuff\n选择-训练-聚合-评估",
        ),
        (
            0.81,
            0.58,
            0.14,
            0.22,
            "#FFF4E8",
            "输出层\nOutput",
            "history.json\nsummary.json\n统计图表",
        ),
    ]

    for x, y, w, h, face, heading, body in boxes:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.012,rounding_size=0.02",
            linewidth=1.6,
            edgecolor="#295F4E",
            facecolor=face,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h * 0.72, heading, ha="center", va="center", fontsize=13, fontweight="bold", color="#173C34")
        ax.text(x + w / 2, y + h * 0.34, body, ha="center", va="center", fontsize=11, color="#2A2A2A")

    for i in range(len(boxes) - 1):
        x1 = boxes[i][0] + boxes[i][2]
        y1 = boxes[i][1] + boxes[i][3] / 2
        x2 = boxes[i + 1][0]
        arrow = FancyArrowPatch(
            (x1 + 0.01, y1),
            (x2 - 0.01, y1),
            arrowstyle="-|>",
            mutation_scale=16,
            linewidth=1.8,
            color="#2D6A4F",
        )
        ax.add_patch(arrow)

    sched = FancyBboxPatch(
        (0.28, 0.36),
        0.44,
        0.12,
        boxstyle="round,pad=0.015,rounding_size=0.02",
        linewidth=1.6,
        linestyle="--",
        edgecolor="#446B9E",
        facecolor="#F7FBFF",
    )
    ax.add_patch(sched)
    ax.text(
        0.5,
        0.42,
        "通信调度器 / 时间模型：将物理接触窗口映射为 FL 训练轮次，\n同时记录 slot 级与 FLOPs/带宽级虚拟时间",
        ha="center",
        va="center",
        fontsize=11.5,
        color="#234A7A",
        fontweight="bold",
    )

    for anchor_x in (0.51, 0.70):
        ax.add_patch(
            FancyArrowPatch(
                (anchor_x, 0.58),
                (0.5, 0.48),
                arrowstyle="-",
                linewidth=1.2,
                color="#446B9E",
            )
        )

    qa = FancyBboxPatch(
        (0.08, 0.08),
        0.84,
        0.18,
        boxstyle="round,pad=0.018,rounding_size=0.02",
        linewidth=1.4,
        edgecolor="#6A7C3B",
        facecolor="#F7FAEE",
    )
    ax.add_patch(qa)
    ax.text(
        0.5,
        0.19,
        "人机协同质控闭环",
        ha="center",
        va="center",
        fontsize=14,
        fontweight="bold",
        color="#46591B",
    )
    ax.text(
        0.5,
        0.11,
        "Prompt 拆解 -> 代码生成/实验执行 -> Ruff 静态检查 -> JSON/图表核查 -> 人工复核结论与边界",
        ha="center",
        va="center",
        fontsize=11.5,
        color="#2F2F2F",
    )

    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build_summary_figure(path: Path, grid_data: list[dict]) -> None:
    configure_matplotlib()
    gs_values = sorted({item["gs_count"] for item in grid_data})
    sat_values = sorted({item["sat_count"] for item in grid_data})

    max_acc = np.zeros((len(gs_values), len(sat_values)))
    final_acc = np.zeros((len(gs_values), len(sat_values)))
    contact_rate = np.zeros((len(gs_values), len(sat_values)))

    lookup = {(item["gs_count"], item["sat_count"]): item for item in grid_data}
    for i, gs in enumerate(gs_values):
        for j, sat in enumerate(sat_values):
            item = lookup[(gs, sat)]
            max_acc[i, j] = item["max_acc"]
            final_acc[i, j] = item["final_acc"]
            contact_rate[i, j] = item["contact_rate"]

    fig = plt.figure(figsize=(15, 10), dpi=220)
    gs = fig.add_gridspec(2, 2, height_ratios=[1, 0.92], width_ratios=[1.05, 1])
    ax1 = fig.add_subplot(gs[0, 0])
    ax2 = fig.add_subplot(gs[0, 1])
    ax3 = fig.add_subplot(gs[1, 0])
    ax4 = fig.add_subplot(gs[1, 1])

    fig.suptitle("SpaceFL 网格实验结果汇总（FedAvg，12 组 GS×SAT 对比）", fontsize=20, fontweight="bold", color="#173C34")

    heat = ax1.imshow(max_acc, cmap="YlGn", vmin=0.35, vmax=0.8, aspect="auto")
    ax1.set_xticks(range(len(sat_values)))
    ax1.set_xticklabels(sat_values)
    ax1.set_yticks(range(len(gs_values)))
    ax1.set_yticklabels(gs_values)
    ax1.set_xlabel("卫星数 SAT")
    ax1.set_ylabel("地面站数 GS")
    ax1.set_title("最高准确率热力图")
    for i in range(len(gs_values)):
        for j in range(len(sat_values)):
            color = "#173C34" if max_acc[i, j] > 0.6 else "#FFFFFF"
            ax1.text(j, i, f"{max_acc[i, j] * 100:.1f}%", ha="center", va="center", fontsize=11, fontweight="bold", color=color)
    fig.colorbar(heat, ax=ax1, fraction=0.046, pad=0.04)

    for sat in sat_values:
        series = [lookup[(gs, sat)]["contact_rate"] * 100 for gs in gs_values]
        ax2.plot(gs_values, series, marker="o", linewidth=2.0, label=f"SAT={sat}")
    ax2.set_title("接触率随地面站数变化")
    ax2.set_xlabel("地面站数 GS")
    ax2.set_ylabel("接触率 / %")
    ax2.grid(True, alpha=0.25)
    ax2.legend(frameon=False)

    sat_focus = 7
    best_series = [lookup[(gs, sat_focus)]["max_acc"] * 100 for gs in gs_values]
    final_series = [lookup[(gs, sat_focus)]["final_acc"] * 100 for gs in gs_values]
    x = np.arange(len(gs_values))
    width = 0.34
    ax3.bar(x - width / 2, best_series, width=width, color="#3A7D44", label="最高准确率")
    ax3.bar(x + width / 2, final_series, width=width, color="#9CC5A1", label="最终准确率")
    ax3.set_xticks(x)
    ax3.set_xticklabels([f"GS={gs}" for gs in gs_values])
    ax3.set_ylim(0, 85)
    ax3.set_ylabel("准确率 / %")
    ax3.set_title("SAT=7 时的峰值-最终值落差")
    ax3.grid(True, axis="y", alpha=0.2)
    ax3.legend(frameon=False)

    ax4.axis("off")
    best_item = max(grid_data, key=lambda item: item["max_acc"])
    best_final = max(grid_data, key=lambda item: item["final_acc"])
    info_lines = [
        "统一控制变量：",
        "Kepler 后端，500 km 高度，53° 倾角，MNIST，",
        "2 类/客户端，local epoch=2，batch=32，lr=0.01。",
        "",
        "当前同口径已完成：",
        "12 组 FedAvg 网格对比，GS∈{2,4,6,8}，SAT∈{3,5,7}。",
        "",
        "关键结果：",
        f"1. 接触率由 6.8% 提升至 24.2%，与 GS 基本线性相关。",
        f"2. 最高准确率最佳为 GS={best_item['gs_count']}, SAT={best_item['sat_count']}，达到 {best_item['max_acc'] * 100:.2f}%。",
        f"3. 最终准确率最佳为 GS={best_final['gs_count']}, SAT={best_final['sat_count']}，达到 {best_final['final_acc'] * 100:.2f}%。",
        "4. 峰值与最终值存在明显落差，说明极端 non-IID 下仍有震荡。",
    ]
    ax4.text(
        0.02,
        0.96,
        "\n".join(info_lines),
        ha="left",
        va="top",
        fontsize=11.5,
        color="#222222",
        bbox={"boxstyle": "round,pad=0.6", "facecolor": "#F7FAEE", "edgecolor": "#6A7C3B"},
    )

    fig.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10.5) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color="1F1F1F")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_col_widths(table, widths_inch: list[float]) -> None:
    for col_idx, width in enumerate(widths_inch):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(width)


def add_experiment_design_table(doc: Document) -> None:
    add_cn_paragraph(
        doc,
        "表：当前报告中可直接支撑结论的实验对比设计。这里明确区分“已完成同口径对比”和“补充案例”，避免跨设置直接横比。",
        bold_prefix="表：",
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["模块", "已完成设置", "统一控制变量", "评价指标", "状态/用途"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)

    rows = [
        [
            "主对比",
            "FedAvg 网格：GS={2,4,6,8}，SAT={3,5,7}",
            "Kepler，500 km，53°，MNIST，2 类/客户端，epoch=2，batch=32，lr=0.01",
            "contact rate，max/final acc，elapsed sec",
            "已完成，可直接支撑主结论",
        ],
        [
            "波动案例",
            "GS=8，SAT=7 作为最优峰值案例",
            "沿用主对比全部设置",
            "min/max/mean/std accuracy，trend curve",
            "用于解释非 IID 震荡",
        ],
        [
            "补充算法案例",
            "FedProx，5GS×10SAT，异构轨道，μ=0.01",
            "MNIST，300 rounds，CappedSelector(max_count=5)",
            "final/max acc，virtual time，oscillation range",
            "可作补充说明，不与主网格直接横比",
        ],
        [
            "待补实验",
            "FedProx/FedBuff 同口径网格；μ 消融；selector/time model 消融",
            "需保持数据划分、轨道与选择器不变",
            "收敛稳定性、峰值-最终值差、通信开销",
            "后续严谨性增强方向",
        ],
    ]

    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)

    set_table_col_widths(table, [0.8, 1.4, 2.0, 1.2, 1.1])


def enhance_docx() -> None:
    grid_data = load_json(GRID_SUMMARY_PATH)
    fedprox_summary = load_json(FEDPROX_SUMMARY_PATH)

    if not BACKUP_PATH.exists():
        shutil.copyfile(SOURCE_DOCX_PATH, BACKUP_PATH)

    doc = Document(SOURCE_DOCX_PATH)

    doc.add_page_break()

    heading = doc.add_paragraph(style="Heading 1")
    run = heading.add_run("3.4 工程质量控制、框架展示与实验严谨性补充")
    set_run_font(run, size=16, bold=True, color="000000")

    add_cn_paragraph(
        doc,
        "这一节用于补齐原报告中相对薄弱的三部分：工程质量门禁、SpaceFL 模型框架可视化，以及实验对比设计的明确说明。原则上只把已经被代码、JSON 和图表验证过的结果写入结论，把尚未同口径完成的实验单独标记为“待补”。",
    )

    sub1 = doc.add_paragraph(style="Heading 2")
    run = sub1.add_run("3.4.1 工程质量门禁：Ruff 不是装饰，而是 AI 代码协作的收口环节")
    set_run_font(run, size=13, bold=True, color="000000")
    add_cn_paragraph(
        doc,
        "项目在 pyproject.toml 中显式配置了 Ruff，目标版本为 Python 3.9，行宽为 100，并统一启用了 F/E/W/I/N/B/SIM/UP/C4/RUF/PERF 等规则集。这意味着 AI 生成或修改的代码并不是“能跑就算完成”，而是必须接受静态规范、潜在 bug、导入排序、命名风格和性能建议的共同约束。",
    )
    add_cn_paragraph(
        doc,
        "现有检查记录 ruff_chk.txt 显示，当前仍有 26 条待处理问题，主要集中在 examples/standard_experiment.py，类型包括 E702（单行多语句）、C401/C408（可简化表达式）和 F841（未使用变量）。这一点恰好体现了人机协同写作中最重要的严谨性：AI 负责提速，但最终交付前必须经过静态门禁和人工复核，不能把“生成出来的代码”直接等同于“可发表级代码”。",
    )
    add_cn_paragraph(
        doc,
        "因此，在报告中提到 Ruff 的意义并不是为了展示工具名，而是为了说明我们采用了“Prompt 生成 -> 代码落地 -> Ruff 检查 -> 实验运行 -> 结果复核”的闭环，这比单纯描述 AI 帮忙写代码更能体现科研工作流的可信度。",
    )

    sub2 = doc.add_paragraph(style="Heading 2")
    run = sub2.add_run("3.4.2 SpaceFL 模型框架展示")
    set_run_font(run, size=13, bold=True, color="000000")
    doc.add_picture(str(FRAMEWORK_PNG), width=Inches(6.2))
    add_caption(doc, "图：SpaceFL 的环境层-轨道层-模拟层-FL 层整体框架，以及用于人机协同的质控闭环。")
    add_cn_paragraph(
        doc,
        "从技术结构看，SpaceFL 并不是单一训练脚本，而是一个分层研究框架：Environment 层提供天体与地面站环境，Orbit 层负责 Kepler 与 Skyfield 双后端轨道力学，Simulator 层把可见性与时隙推进转换为接触矩阵，FL 层再在此基础上运行 FedAvg、FedProx 和 FedBuff。CommunicationScheduler 与 TimeModel 位于模拟层与 FL 层之间，负责把“物理上何时能通信”映射为“训练上何时能推进一轮”。这正是 SpaceFL 相较于普通 FL 仿真代码的核心研究价值所在。",
    )

    sub3 = doc.add_paragraph(style="Heading 2")
    run = sub3.add_run("3.4.3 实验对比设计与结果汇总")
    set_run_font(run, size=13, bold=True, color="000000")
    add_experiment_design_table(doc)
    doc.add_picture(str(SUMMARY_PNG), width=Inches(6.2))
    add_caption(doc, "图：12 组 FedAvg 网格实验的中文汇总。左上为最高准确率热力图，右上为接触率随 GS 的变化，左下为 SAT=7 时峰值与最终值的落差。")

    best_item = max(grid_data, key=lambda item: item["max_acc"])
    best_final = max(grid_data, key=lambda item: item["final_acc"])
    add_cn_paragraph(
        doc,
        f"从已经完成的 12 组同口径 FedAvg 网格实验看，最稳妥的结论有三条。第一，接触率几乎只随地面站数增加而线性提升：GS 从 2 提升到 8 时，接触率由 6.8% 提高到 24.2%。第二，峰值准确率最佳配置为 GS={best_item['gs_count']}、SAT={best_item['sat_count']}，最高达到 {best_item['max_acc'] * 100:.2f}%；但最终准确率最佳配置其实是 GS={best_final['gs_count']}、SAT={best_final['sat_count']}，为 {best_final['final_acc'] * 100:.2f}%。第三，峰值与最终值之间存在显著落差，说明极端 non-IID 条件下的客户端漂移和单客户端主导问题仍未被完全解决。",
    )

    fedprox_result = fedprox_summary["results_summary"]
    add_cn_paragraph(
        doc,
        f"此外，results/experiment_summary.json 中还保留了一组 FedProx 补充案例：在 5GS×10SAT、异构轨道、μ=0.01 的设置下，模型最高准确率可到 {fedprox_result['max_accuracy'] * 100:.2f}%，但最终准确率回落到 {fedprox_result['final_accuracy'] * 100:.2f}%。由于这组实验与前述 12 组 FedAvg 网格并不同口径，它更适合作为“方法可行但还需严格对齐基线”的补充证据，而不应被直接拿来横向宣称优于 FedAvg。",
    )

    sub4 = doc.add_paragraph(style="Heading 2")
    run = sub4.add_run("3.4.4 人机协同下的严谨性提升方向")
    set_run_font(run, size=13, bold=True, color="000000")
    add_cn_paragraph(
        doc,
        "如果希望这份报告在人机协同和实验严谨性上更像一篇高分作业，而不是“AI 帮我整理了一遍”，后续最值得补的不是更花哨的文字，而是三类同口径实验：其一，固定数据划分、轨道参数与 selector 后，对 FedAvg / FedProx / FedBuff 做真正的一一对照；其二，对 μ、参与客户端数和 time model 做消融，验证震荡是否被稳定抑制；其三，报告里明确区分“代码已实现”“实验已跑通”“结论已被同口径验证”三个层级。这样写，既诚实，也更符合科研报告的证据规范。",
    )

    doc.save(OUTPUT_DOCX_PATH)


def main() -> None:
    ensure_dir(GENERATED_DIR)
    grid_data = load_json(GRID_SUMMARY_PATH)
    build_framework_figure(FRAMEWORK_PNG)
    build_summary_figure(SUMMARY_PNG, grid_data)
    enhance_docx()
    print(f"updated_docx={OUTPUT_DOCX_PATH}")
    print(f"framework_png={FRAMEWORK_PNG}")
    print(f"summary_png={SUMMARY_PNG}")


if __name__ == "__main__":
    main()
